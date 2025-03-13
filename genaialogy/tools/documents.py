"""Tools for processing and analyzing documents."""

from pathlib import Path
import re
import subprocess

from docx import Document as DocxDocument
from llama_index.core import Document, VectorStoreIndex
from pypdf import PdfReader
from striprtf.striprtf import rtf_to_text

DEFAULT_FILE_EXTENSIONS = [".pdf", ".docx", ".doc", ".txt", ".rtf"]


class DocumentLoader:
    """Load documents from a directory into LlamaIndex."""

    def __init__(
        self, directory, text_output_dir=None, file_extensions=DEFAULT_FILE_EXTENSIONS
    ):
        self.documents = None
        self.text_output_dir = text_output_dir
        self.file_extensions = file_extensions
        self.directory = directory

    def to_index(self, storage_dir="index_storage"):
        """Load documents into LlamaIndex."""
        if not self.documents:
            self.from_directory()
        index = VectorStoreIndex.from_documents(self.documents)
        index.storage_context.persist(storage_dir)
        return index

    def load(self):
        """Load documents from a directory into LlamaIndex as Documents."""

        documents = []

        if self.text_output_dir:
            Path(self.text_output_dir).mkdir(parents=True, exist_ok=True)

        nloaded = 0
        for file_path in sorted(Path(self.directory).glob("**/*")):
            if not file_path.is_file():
                continue
            if self.file_extensions and not any(
                filter.lower() in str(file_path).lower()
                for filter in self.file_extensions
            ):
                print(f"Skipping file: {file_path}")
                continue
            print(f"Loading document: {file_path}")
            if file_path.suffix.lower() == ".pdf":
                text = self.extract_text_from_pdf(file_path)
            elif file_path.suffix.lower() in (".doc", ".docx"):
                text = self.extract_text_from_word(file_path)
            elif file_path.suffix.lower() == ".txt":
                text = self.extract_text_from_txt(file_path)
            elif file_path.suffix.lower() == ".rtf":
                text = self.extract_text_from_rtf(file_path)
            else:
                print(f"Skipping unsupported file type: {file_path}")
                continue

            if self.is_text_empty(text):
                print(f"No text extracted - skipping: {file_path}")
                continue

            if self.text_output_dir:
                output_path = (
                    Path(self.text_output_dir) / file_path.stem
                )  # Strip extension
                with open(output_path.with_suffix(".txt"), "w", encoding="utf-8") as f:
                    print(f"Saving text to: {output_path}.txt")
                    f.write(text)

            print(f"Character count: {len(text)}")
            print(f"Word count: {self.word_count(text)}")

            documents.append(
                Document(
                    text=text,
                    metadata={
                        "file_name": file_path.name,
                    },
                )
            )
            nloaded += 1
            print(120 * "=")

        self.documents = documents

    @classmethod
    def is_text_empty(text):
        """
        Check if text is empty after removing all non-printable characters.
        """
        cleaned_text = re.sub(r"\s+", "", text)  # Remove all whitespace characters
        cleaned_text = "".join(
            c for c in cleaned_text if c.isprintable()
        )  # Remove non-printable chars
        return len(cleaned_text) == 0

    @classmethod
    def word_count(text):
        return len(re.findall(r"\b\w+\b", text))

    @classmethod
    def extract_text_from_pdf(pdf_path):
        """
        Extract text from a PDF file.
        """
        try:
            with open(pdf_path, "rb") as f:
                reader = PdfReader(f)
                return "\n".join([page.extract_text() or "" for page in reader.pages])
        except Exception as e:
            print(f"Failed to extract text from {pdf_path}: {e}")
            return ""

    @classmethod
    def extract_text_from_docx(doc_path):
        """
        Extract text from a .docx file.
        """
        try:
            doc = DocxDocument(doc_path)
            return "\n".join([para.text for para in doc.paragraphs])
        except Exception as e:
            print(f"Failed to extract text from {doc_path}: {e}")
            return ""

    @classmethod
    def extract_text_from_doc(doc_path):
        """
        Extract text from a .doc file using LibreOffice in Google Colab.
        """
        try:
            output_txt_path = Path(doc_path).with_suffix(".txt")
            subprocess.run(
                [
                    "soffice",
                    "--headless",
                    "--convert-to",
                    "txt:Text",
                    "--outdir",
                    str(output_txt_path.parent),
                    str(doc_path),
                ],
                check=True,
            )
            with open(output_txt_path, "r", encoding="utf-8") as f:
                text = f.read()
            output_txt_path.unlink()  # Remove the temporary .txt file
            return text
        except Exception as e:
            print(f"Failed to extract text from {doc_path}: {e}")
            return ""

    @classmethod
    def extract_text_from_word(doc_path):
        """
        Extract text from a Word file (.doc or .docx).
        Uses python-docx for .docx and pandoc for .doc.
        """
        if doc_path.suffix.lower() == ".docx":
            return DocumentLoader.extract_text_from_docx(doc_path)
        elif doc_path.suffix.lower() == ".doc":  # Convert .doc to text using pandoc
            return DocumentLoader.extract_text_from_doc(doc_path)
        else:
            print(f"Unsupported file type: {doc_path}")
            return ""

    @classmethod
    def extract_text_from_txt(txt_path):
        """
        Extract text from a plain text (.txt) file.
        """
        try:
            with open(txt_path, "r", encoding="utf-8") as f:
                return f.read()
        except Exception as e:
            print(f"Failed to extract text from {txt_path}: {e}")
            return ""

    @classmethod
    def extract_text_from_rtf(rtf_path):
        """
        Extract text from an RTF file.

        :param rtf_path: Path to the .rtf file.
        :return: Extracted plain text.
        """
        try:
            with open(rtf_path, "r", encoding="utf-8") as f:
                rtf_content = f.read()
            return rtf_to_text(rtf_content)
        except Exception as e:
            print(f"Failed to extract text from {rtf_path}: {e}")
            return ""
